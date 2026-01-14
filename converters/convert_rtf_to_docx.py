#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Модуль для конвертации RTF файлов в DOCX формат.
Использует word_reader для чтения RTF и python-docx для создания DOCX документов.
"""

import sys
from pathlib import Path
from typing import Optional, List

try:
    from converters.word_reader import read_rtf_blocks, read_rtf, ReaderConfig, TextBlock
    WORD_READER_AVAILABLE = True
except ImportError:
    WORD_READER_AVAILABLE = False
    # Fallback на прямое использование striprtf
    try:
        import striprtf.striprtf as rtf
        RTF_AVAILABLE = True
    except ImportError:
        RTF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ConversionError(Exception):
    """Базовый класс для ошибок конвертации."""
    pass


def read_rtf_content(file_path: Path, encodings: List[str] = None) -> str:
    """
    Читает содержимое RTF файла с попыткой разных кодировок.
    
    Args:
        file_path: Путь к RTF файлу
        encodings: Список кодировок для попытки чтения
        
    Returns:
        Содержимое RTF файла как строка
        
    Raises:
        ConversionError: Если не удалось прочитать файл
    """
    if encodings is None:
        encodings = ['utf-8', 'cp1251', 'windows-1251', 'latin-1']
    
    for encoding in encodings:
        try:
            content = file_path.read_text(encoding=encoding, errors='ignore')
            if content:
                return content
        except Exception:
            continue
    
    raise ConversionError(f"Не удалось прочитать RTF файл {file_path} с доступными кодировками")


def extract_text_from_rtf(rtf_content: str) -> str:
    """
    Извлекает текст из RTF содержимого (fallback метод).
    
    Args:
        rtf_content: Содержимое RTF файла
        
    Returns:
        Извлеченный текст
        
    Raises:
        ConversionError: Если не удалось извлечь текст
    """
    if not RTF_AVAILABLE:
        raise ConversionError(
            "Библиотека striprtf не установлена. "
            "Установите её командой: pip install striprtf"
        )
    
    try:
        plain_text = rtf.rtf_to_text(rtf_content)
        return plain_text or ""
    except Exception as e:
        raise ConversionError(f"Ошибка при извлечении текста из RTF: {e}") from e


def convert_rtf_to_docx(
    input_file: Path,
    output_file: Optional[Path] = None,
    preserve_formatting: bool = True
) -> Path:
    """
    Конвертирует RTF файл в DOCX формат.
    
    Args:
        input_file: Путь к входному RTF файлу
        output_file: Путь к выходному DOCX файлу (если None, создается рядом с входным)
        preserve_formatting: Сохранять ли форматирование (разбиение на параграфы)
        
    Returns:
        Путь к созданному DOCX файлу
        
    Raises:
        ConversionError: При ошибках конвертации
    """
    if not DOCX_AVAILABLE:
        raise ConversionError(
            "Библиотека python-docx не установлена. "
            "Установите её командой: pip install python-docx"
        )
    
    if not input_file.exists():
        raise ConversionError(f"Файл не найден: {input_file}")
    
    if input_file.suffix.lower() != '.rtf':
        raise ConversionError(f"Файл должен быть в формате .rtf, получен: {input_file.suffix}")
    
    # Определяем выходной файл
    if output_file is None:
        output_file = input_file.with_suffix('.docx')
    else:
        output_file = Path(output_file)
        if output_file.suffix.lower() != '.docx':
            output_file = output_file.with_suffix('.docx')
    
    # Читаем RTF используя word_reader если доступен
    if WORD_READER_AVAILABLE:
        try:
            # Используем word_reader для более качественного чтения
            config = ReaderConfig(clean=False)  # Не очищаем, чтобы сохранить структуру
            blocks = read_rtf_blocks(input_file, config)
            
            if not blocks:
                raise ConversionError("RTF файл не содержит текста")
            
            # Создаем DOCX документ
            doc = Document()
            
            if preserve_formatting:
                # Используем блоки для сохранения структуры
                for block in blocks:
                    if block.text.strip():
                        doc.add_paragraph(block.text)
                    else:
                        # Пустые блоки - добавляем пустой параграф
                        doc.add_paragraph()
            else:
                # Объединяем все блоки в один параграф
                all_text = "\n".join(block.text for block in blocks)
                doc.add_paragraph(all_text)
                
        except Exception as e:
            # Fallback на старый метод
            rtf_content = read_rtf_content(input_file)
            plain_text = extract_text_from_rtf(rtf_content)
            
            if not plain_text:
                raise ConversionError("RTF файл не содержит текста")
            
            doc = Document()
            
            if preserve_formatting:
                lines = plain_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
                    else:
                        doc.add_paragraph()
            else:
                doc.add_paragraph(plain_text)
    else:
        # Используем старый метод без word_reader
        rtf_content = read_rtf_content(input_file)
        plain_text = extract_text_from_rtf(rtf_content)
        
        if not plain_text:
            raise ConversionError("RTF файл не содержит текста")
        
        doc = Document()
        
        if preserve_formatting:
            lines = plain_text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()
        else:
            doc.add_paragraph(plain_text)
    
    # Сохраняем документ
    try:
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))
        return output_file
    except Exception as e:
        raise ConversionError(f"Ошибка при создании DOCX файла: {e}") from e


def convert_directory(
    input_dir: Path,
    output_dir: Optional[Path] = None,
    preserve_formatting: bool = True,
    recursive: bool = False
) -> List[Path]:
    """
    Конвертирует все RTF файлы в директории в DOCX.
    
    Args:
        input_dir: Директория с RTF файлами
        output_dir: Директория для сохранения DOCX файлов (если None, создается рядом)
        preserve_formatting: Сохранять ли форматирование
        recursive: Обрабатывать ли поддиректории рекурсивно
        
    Returns:
        Список путей к созданным DOCX файлам
    """
    if not input_dir.exists() or not input_dir.is_dir():
        raise ConversionError(f"Директория не найдена: {input_dir}")
    
    # Определяем выходную директорию
    if output_dir is None:
        script_dir = Path(__file__).parent.absolute()
        output_dir = script_dir / 'docx_output'
    else:
        output_dir = Path(output_dir)
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Находим все RTF файлы
    if recursive:
        rtf_files = list(input_dir.rglob('*.rtf'))
    else:
        rtf_files = list(input_dir.glob('*.rtf'))
    
    if not rtf_files:
        raise ConversionError(f"В директории {input_dir} не найдено RTF файлов")
    
    converted_files = []
    errors = []
    
    for rtf_file in rtf_files:
        try:
            # Определяем относительный путь для сохранения структуры
            if recursive:
                relative_path = rtf_file.relative_to(input_dir)
                output_file = output_dir / relative_path.with_suffix('.docx')
            else:
                output_file = output_dir / rtf_file.with_suffix('.docx').name
            
            # Создаем поддиректории при необходимости
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Конвертируем
            result = convert_rtf_to_docx(rtf_file, output_file, preserve_formatting)
            converted_files.append(result)
            print(f"✓ {rtf_file.name} -> {result.name}")
            
        except Exception as e:
            error_msg = f"✗ Ошибка при конвертации {rtf_file.name}: {e}"
            print(error_msg)
            errors.append((rtf_file, str(e)))
    
    if errors:
        print(f"\n⚠ Ошибок при конвертации: {len(errors)}")
    
    return converted_files


def main():
    """Основная функция для CLI."""
    import argparse
    
    # Определяем текущую директорию скрипта
    script_dir = Path(__file__).parent.absolute()
    default_input_dir = script_dir / 'words_input'
    default_output_dir = script_dir / 'docx_output'
    
    parser = argparse.ArgumentParser(
        description='Конвертация RTF файлов в DOCX формат',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
Примеры использования:
  python converters/convert_rtf_to_docx.py                    # Обработает {default_input_dir.name}/
  python converters/convert_rtf_to_docx.py document.rtf
  python converters/convert_rtf_to_docx.py document.rtf -o output.docx
  python converters/convert_rtf_to_docx.py words_input/ --output-dir docx_output/
  python converters/convert_rtf_to_docx.py words_input/ -r  # рекурсивно
        """
    )
    
    parser.add_argument(
        'input_path',
        type=str,
        nargs='?',
        default=str(default_input_dir),
        help=f'Путь к RTF файлу или директории с RTF файлами (по умолчанию: {default_input_dir.name}/)'
    )
    parser.add_argument(
        '-o', '--output',
        type=str,
        default=None,
        help='Путь к выходному DOCX файлу (для одного файла)'
    )
    parser.add_argument(
        '--output-dir',
        type=str,
        default=None,
        help='Директория для сохранения DOCX файлов (для директории)'
    )
    parser.add_argument(
        '--no-formatting',
        action='store_true',
        help='Не сохранять форматирование (весь текст в один параграф)'
    )
    parser.add_argument(
        '-r', '--recursive',
        action='store_true',
        help='Обрабатывать поддиректории рекурсивно'
    )
    
    args = parser.parse_args()
    
    # Преобразуем относительный путь в абсолютный
    input_path = Path(args.input_path)
    if not input_path.is_absolute():
        input_path = script_dir / input_path
    
    if not input_path.exists():
        print(f"Ошибка: путь не существует: {input_path}")
        sys.exit(1)
    
    preserve_formatting = not args.no_formatting
    
    try:
        if input_path.is_file():
            # Конвертация одного файла
            output_path = Path(args.output) if args.output else None
            if output_path and not output_path.is_absolute():
                output_path = script_dir / output_path
            
            print(f"Конвертация: {input_path.name}")
            result = convert_rtf_to_docx(input_path, output_path, preserve_formatting)
            print(f"✓ Успешно сохранено: {result}")
            
        elif input_path.is_dir():
            # Конвертация директории
            output_dir = Path(args.output_dir) if args.output_dir else None
            if output_dir and not output_dir.is_absolute():
                output_dir = script_dir / output_dir
            
            print(f"Обработка директории: {input_path}")
            converted = convert_directory(
                input_path,
                output_dir,
                preserve_formatting,
                args.recursive
            )
            print(f"\n✓ Конвертировано файлов: {len(converted)}")
            print(f"Результаты сохранены в: {output_dir or (input_path / 'docx_output')}")
            
        else:
            print(f"Ошибка: указанный путь не является файлом или директорией")
            sys.exit(1)
            
    except ConversionError as e:
        print(f"Ошибка конвертации: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Неожиданная ошибка: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()

